"""DOCX generation engine.

Design goals:
  * Never recreate the document. Open the master template and replace ONLY
    placeholders, preserving fonts, styles, tables, margins, headers, footers.
  * Handle placeholders that Word splits across multiple runs.
  * Optional signature suppression for the neutral ``No signature`` role.
"""
from __future__ import annotations

import logging
import re
from typing import Dict, Iterable, List

from docx import Document

from notice_app.services.placeholder_map import (
    ALIAS_PLACEHOLDERS,
    SENDER_ROLE_TOKEN,
    SENDER_NAME_TOKEN,
    RECORD_PLACEHOLDERS,
    token,
)

logger = logging.getLogger(__name__)

_TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def _substitute(text: str, mapping: Dict[str, str]) -> str:
    """Replace every {{token}} in a plain string (case-insensitive keys)."""
    def _sub(match: re.Match) -> str:
        key = match.group(1).strip().lower()
        return mapping.get(key, match.group(0))

    return _TOKEN_RE.sub(_sub, text)


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    """Replace {{tokens}} inside a paragraph, tolerant of run-splitting.

    Word frequently breaks a placeholder across several runs (e.g. ``{{``,
    ``sender_name``, ``}}``), which is exactly why token replacement kept
    failing for the bold/underlined signatory lines. Rather than trying to
    edit individual runs, we replace across the complete paragraph:
    take the paragraph's FULL text, substitute across the whole string, then
    write the result back so the placeholder is guaranteed to be replaced no
    matter how Word fragmented it.

    Formatting: the first run's properties are kept and applied to the whole
    line (signatory/placeholder lines are uniformly formatted in the template,
    so this preserves the intended look). The caller re-applies any special
    font afterwards (e.g. the bank line).
    """
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    new_text = _substitute(full_text, mapping)
    if new_text == full_text:
        return

    runs = paragraph.runs
    if runs:
        # Put all replaced text into the first run; clear the rest. This keeps
        # the first run's formatting and collapses any split placeholder.
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
    else:
        # No runs exposed (text lives in nested wrappers). Fall back to the
        # python-docx text setter, which rebuilds the paragraph's runs.
        paragraph.text = new_text


def _paragraph_has_token(paragraph, name: str) -> bool:
    # Use paragraph.text (full content) rather than only paragraph.runs, so
    # tokens Word nested in wrappers are still detected.
    return token(name).lower() in paragraph.text.lower()


def _set_paragraph_font(paragraph, font_name: str, size_pt: int) -> None:
    """Force a font name + size on every run of a paragraph (run + rPr level)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)


def _delete_paragraph(paragraph) -> None:
    """Remove a paragraph element entirely (no leftover blank line)."""
    element = paragraph._element
    element.getparent().remove(element)
    element._p = element._element = None


def _iter_all_paragraphs(document) -> Iterable:
    """Yield paragraphs from body, tables, headers and footers (all sections)."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def build_mapping(record: dict, sender_name: str, sender_role: str,
                  unsigned_role: str, date_value: str = "",
                  subject_value: str = "") -> Dict[str, str]:
    """Construct the token->value mapping for a single notice."""
    mapping: Dict[str, str] = {}
    for token_name, attr in RECORD_PLACEHOLDERS.items():
        value = record.get(attr, "")
        mapping[token_name] = "" if value is None else str(value)

    # Alias tokens (e.g. appno -> acknowledgement_no, bank_name -> bank).
    for alias, attr in ALIAS_PLACEHOLDERS.items():
        value = record.get(attr, "")
        mapping[alias] = "" if value is None else str(value)

    # Computed tokens.
    mapping["date"] = date_value or ""
    mapping["subject"] = subject_value or ""

    mapping[SENDER_ROLE_TOKEN] = sender_role or ""
    if sender_role == unsigned_role:
        # The optional sender name is suppressed entirely.
        mapping[SENDER_NAME_TOKEN] = ""
    else:
        mapping[SENDER_NAME_TOKEN] = sender_name or ""
    return mapping


def render_document(template_path: str, record: dict, sender_name: str,
                    sender_role: str, unsigned_role: str,
                    date_value: str = "", subject_value: str = ""):
    """Open the template, apply replacements, return a python-docx Document.

    Placeholder replacement is delegated to docxtpl (Jinja), the same proven
    template engine. docxtpl reliably substitutes
    {{tokens}} no matter how Word split them across runs, which the previous
    hand-rolled python-docx logic failed to do for the signatory lines.

    After rendering, the underlying python-docx Document is post-processed for
    the behaviours docxtpl does not cover: optional signature suppression and
    the {{account_table}} expansion.
    """
    from docxtpl import DocxTemplate

    mapping = build_mapping(
        record, sender_name, sender_role, unsigned_role,
        date_value=date_value, subject_value=subject_value,
    )

    # docxtpl context: every token maps to its value. The account table is kept
    # as a literal marker so the python-docx post-step can expand it into a
    # real bordered table.
    context = dict(mapping)
    context["account_table"] = token("account_table")

    tpl = DocxTemplate(template_path)
    try:
        tpl.render(context)
    except Exception:  # noqa: BLE001 - never fail generation on a render error.
        logger.exception("docxtpl render failed for %s; falling back", template_path)
        return _render_document_fallback(
            template_path, record, sender_name, sender_role,
            unsigned_role, mapping,
        )

    document = tpl.docx  # underlying python-docx Document.

    suppress_name = sender_role == unsigned_role

    if suppress_name:
        _suppress_empty_sender_line(document, sender_role)

    # Expand the {{account_table}} token into a real bordered table.
    _insert_account_table(document, record)

    _log_unreplaced(document, template_path)
    return document


def _suppress_empty_sender_line(document, sender_role: str) -> None:
    """Remove the sender-name paragraph left blank for an unsigned notice.

    After docxtpl renders {{sender_name}} to an empty string, the paragraph
    that held only the sender name is now empty (or whitespace). We delete it
    so the sender role moves up with no blank line. A paragraph that also
    carries the sender role text is never deleted.
    """
    role = (sender_role or "").strip()
    for paragraph in list(_iter_all_paragraphs(document)):
        text = paragraph.text
        if text.strip():
            continue
        # Empty paragraph: delete only if it is NOT the sender role line and
        # sits directly above/below content (typical signatory block).
        # We simply drop empty paragraphs adjacent to the sender role.
        _maybe_delete_if_near_sender_role(document, paragraph, role)


def _maybe_delete_if_near_sender_role(document, paragraph, role: str) -> None:
    """Delete an empty paragraph if the next non-empty line is the sender role.

    This mirrors the old behaviour: only the blank left by the suppressed
    sender name (which sits immediately before the sender role) is removed,
    not every blank line in the document.
    """
    if not role:
        return
    body_paras = document.paragraphs
    try:
        idx = body_paras.index(paragraph)
    except ValueError:
        return
    # Look at the next paragraph; if it is the sender role line, drop this blank.
    if idx + 1 < len(body_paras):
        nxt = body_paras[idx + 1].text.strip().lower()
        if nxt == role.lower():
            _delete_paragraph(paragraph)


def _render_document_fallback(template_path, record, sender_name, sender_role,
                              unsigned_role, mapping):
    """Pure python-docx replacement, used only if docxtpl rendering fails."""
    document = Document(template_path)
    suppress_name = sender_role == unsigned_role
    for paragraph in list(_iter_all_paragraphs(document)):
        if suppress_name and _paragraph_has_token(paragraph, SENDER_NAME_TOKEN):
            has_sender_role = _paragraph_has_token(paragraph, SENDER_ROLE_TOKEN)
            _replace_in_paragraph(paragraph, mapping)
            if not has_sender_role and not paragraph.text.strip():
                _delete_paragraph(paragraph)
                continue
        else:
            _replace_in_paragraph(paragraph, mapping)
    _insert_account_table(document, record)
    _log_unreplaced(document, template_path)
    return document


def _insert_account_table(document, record: dict) -> None:
    """Replace a {{account_table}} placeholder paragraph with a real table.

    The notice lists this record's transaction row. If no placeholder is
    present, nothing happens.
    """
    target = None
    for paragraph in _iter_all_paragraphs(document):
        if "{{account_table}}" in paragraph.text.lower().replace(" ", ""):
            target = paragraph
            break
        if "{{account_table}}" in paragraph.text.lower():
            target = paragraph
            break
    if target is None:
        return

    # Clear the placeholder text in place.
    if target.runs:
        target.runs[0].text = ""
        for run in target.runs[1:]:
            run.text = ""

    headers = ["Account", "IFSC", "Transaction ID", "Date", "Amount"]
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:  # noqa: BLE001 - style may not exist in the template.
        pass

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    _mark_header_row(table.rows[0])

    row_cells = table.add_row().cells
    row_cells[0].text = str(record.get("account_no", "") or "")
    row_cells[1].text = str(record.get("ifsc", "") or "")
    row_cells[2].text = str(record.get("transaction_id", "") or "")
    row_cells[3].text = str(record.get("transaction_date", "") or "")
    row_cells[4].text = str(record.get("transaction_amount", "") or "")

    _set_table_geometry(table, [1944, 1800, 2232, 1656, 1728])

    # Compact business typography across the whole table (headers stay bold).
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_font(paragraph, "Aptos", 9)

    _apply_table_borders(table)

    # Move the table to immediately after the placeholder paragraph, then remove
    # the now-empty placeholder paragraph so no blank line is left where the
    # {{account_table}} token used to be.
    moved = False
    try:
        tbl_element = table._tbl
        tbl_element.getparent().remove(tbl_element)
        target._p.addnext(tbl_element)
        moved = True
    except Exception:  # noqa: BLE001 - leave at end if the move fails.
        moved = False

    if moved:
        # Delete the emptied placeholder paragraph (the table now occupies its
        # position). Guard against the element already being detached.
        try:
            _delete_paragraph(target)
        except Exception:  # noqa: BLE001
            pass

    # add_table() appended the table at the end of the body, which also leaves
    # trailing empty paragraphs behind. Strip trailing blank paragraphs so the
    # document does not spill onto an extra page.
    _strip_trailing_empty_paragraphs(document)


def _strip_trailing_empty_paragraphs(document) -> None:
    """Remove empty paragraphs at the very end of the document body.

    Word keeps at least one paragraph after the final table, so we stop as soon
    as we hit a non-empty paragraph or a paragraph that is not the last body
    element. This clears the blank line(s) that push content onto page 2.
    """
    paragraphs = document.paragraphs
    for paragraph in reversed(paragraphs):
        if paragraph.text.strip():
            break
        # Only strip paragraphs that are genuinely trailing (no following
        # sibling element such as a table).
        element = paragraph._element
        if element.getnext() is not None:
            break
        _delete_paragraph(paragraph)


def _apply_table_borders(table) -> None:
    """Apply quiet gray borders and a restrained header fill."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            if row_index == 0:
                shade = tc_pr.find(qn("w:shd"))
                if shade is None:
                    shade = OxmlElement("w:shd")
                    tc_pr.append(shade)
                shade.set(qn("w:fill"), "EEF2F6")
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "bottom", "left", "right"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    element.set(qn("w:val"), "single")
                    element.set(qn("w:sz"), "4")
                    element.set(qn("w:space"), "0")
                    element.set(qn("w:color"), "CBD5E1")
                    borders.append(element)


def _set_table_geometry(table, widths_dxa) -> None:
    """Set fixed, explicit DXA geometry for predictable Word rendering."""
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for edge, value in (
                ("top", 80), ("bottom", 80), ("start", 120), ("end", 120)
            ):
                node = tc_mar.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_mar.append(node)
                node.set(qn("w:type"), "dxa")
                node.set(qn("w:w"), str(value))


def _mark_header_row(row) -> None:
    """Mark the first row as a semantic/repeating table header."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _log_unreplaced(document, template_path: str) -> None:
    leftovers: List[str] = []
    for paragraph in _iter_all_paragraphs(document):
        for match in _TOKEN_RE.finditer(paragraph.text):
            leftovers.append(match.group(0))
    if leftovers:
        logger.warning(
            "Unreplaced placeholders in %s: %s",
            template_path,
            ", ".join(sorted(set(leftovers))),
        )


def extract_preview_lines(template_path: str, record: dict, sender_name: str,
                          sender_role: str, unsigned_role: str,
                          date_value: str = "", subject_value: str = ""):
    """Render to an in-memory document and return plain-text lines for preview.

    The web preview mirrors the DOCX content/structure; exact pixel rendering
    would require LibreOffice headless conversion (optional, off by default).
    """
    document = render_document(
        template_path, record, sender_name, sender_role, unsigned_role,
        date_value=date_value, subject_value=subject_value,
    )
    lines = []
    for paragraph in document.paragraphs:
        lines.append({
            "text": paragraph.text,
            "align": str(paragraph.alignment) if paragraph.alignment else "",
            "style": paragraph.style.name if paragraph.style else "",
        })
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return {"paragraphs": lines, "tables": tables}
