"""Report and export generation for the Investigation Intelligence Engine.

Produces Markdown plus lightweight JSON/PDF/DOCX exports from the same local
SQLite intelligence store. PDF/DOCX imports are lazy so the engine still boots
if optional offline packages are missing.
"""
from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from investigation_app.extensions import get_connection
from investigation_app.services import timeline_service


def _case_row(conn, case_uid: str):
    return conn.execute("SELECT * FROM cases WHERE uid = ?", (case_uid,)).fetchone()


def _generated() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _rows(conn, sql: str, args: tuple) -> List[Dict[str, Any]]:
    return [dict(r) for r in conn.execute(sql, args).fetchall()]


def _safe_json(raw: str, default):
    try:
        return json.loads(raw or "")
    except (TypeError, ValueError):
        return default


def build_export_data(case_uid: str, kind: str = "investigation") -> Optional[Dict[str, Any]]:
    """Return the structured export payload, or None if the case is missing."""
    conn = get_connection()
    try:
        case = _case_row(conn, case_uid)
        if case is None:
            return None
        case_id = case["id"]
        evidence = _rows(
            conn,
            "SELECT id, original_name, mime, size, sha256, status, created_at, intel_json "
            "FROM evidence WHERE case_id = ? ORDER BY created_at",
            (case_id,),
        )
        for ev in evidence:
            ev["intelligence"] = _safe_json(ev.pop("intel_json", None), {})
        entities = _rows(
            conn,
            "SELECT e.id, e.type, e.value, e.norm, COUNT(l.evidence_id) AS links "
            "FROM entities e LEFT JOIN entity_links l ON l.entity_id = e.id "
            "WHERE e.case_id = ? GROUP BY e.id ORDER BY links DESC, e.type",
            (case_id,),
        )
        transactions = _rows(
            conn,
            "SELECT t.*, e.original_name FROM transactions t "
            "LEFT JOIN evidence e ON e.id = t.evidence_id "
            "WHERE t.case_id = ? ORDER BY COALESCE(t.layer, 999), t.txn_date, t.id",
            (case_id,),
        )
        messages = _rows(
            conn,
            "SELECT m.*, e.original_name FROM communications m "
            "LEFT JOIN evidence e ON e.id = m.evidence_id "
            "WHERE m.case_id = ? ORDER BY COALESCE(m.timestamp, ''), m.id",
            (case_id,),
        )
        for msg in messages:
            msg["risk_flags"] = _safe_json(msg.get("risk_flags_json"), [])
            msg["urls"] = _safe_json(msg.get("urls_json"), [])
        social_profiles = _rows(
            conn,
            "SELECT sp.*, e.original_name FROM social_profiles sp "
            "LEFT JOIN evidence e ON e.id = sp.evidence_id "
            "WHERE sp.case_id = ? ORDER BY sp.platform, sp.username",
            (case_id,),
        )
        for sp in social_profiles:
            sp["metadata"] = _safe_json(sp.get("metadata_json"), {})
        technical_indicators = _rows(
            conn,
            "SELECT ti.*, e.original_name FROM technical_indicators ti "
            "LEFT JOIN evidence e ON e.id = ti.evidence_id "
            "WHERE ti.case_id = ? ORDER BY ti.type, ti.value",
            (case_id,),
        )
        for ti in technical_indicators:
            ti["metadata"] = _safe_json(ti.get("metadata_json"), {})
        similarities = _rows(
            conn,
            "SELECT s.*, ea.original_name AS a_name, eb.original_name AS b_name "
            "FROM evidence_similarity s "
            "LEFT JOIN evidence ea ON ea.id = s.a_id "
            "LEFT JOIN evidence eb ON eb.id = s.b_id "
            "WHERE s.case_id = ? ORDER BY s.score DESC",
            (case_id,),
        )
        for sim in similarities:
            sim["reasons"] = _safe_json(sim.get("reasons"), [])
        relationships = _rows(
            conn,
            "SELECT r.src_id, r.dst_id, r.weight, s.type AS src_type, s.value AS src_value, "
            "d.type AS dst_type, d.value AS dst_value "
            "FROM relationships r "
            "JOIN entities s ON s.id = r.src_id "
            "JOIN entities d ON d.id = r.dst_id "
            "WHERE r.case_id = ? ORDER BY r.weight DESC LIMIT 200",
            (case_id,),
        )
        failed_stages = _rows(
            conn,
            "SELECT es.evidence_id, e.original_name, es.stage, es.state, es.detail, es.at "
            "FROM evidence_stages es JOIN evidence e ON e.id = es.evidence_id "
            "WHERE e.case_id = ? AND es.state != 'ok' ORDER BY es.at DESC",
            (case_id,),
        )
        timeline = timeline_service.list_events(case_uid)
    finally:
        conn.close()

    leads = _build_leads(transactions, similarities, evidence, messages, social_profiles, technical_indicators)
    return {
        "case": {
            "uid": case["uid"],
            "title": case["title"],
            "reference_no": case["reference_no"],
            "status": case["status"],
            "ai_mode": case["ai_mode"],
            "generated_at": _generated(),
            "report_kind": kind,
        },
        "evidence": evidence,
        "entities": entities,
        "transactions": transactions,
        "messages": messages,
        "social_profiles": social_profiles,
        "technical_indicators": technical_indicators,
        "timeline": timeline,
        "similar_evidence": similarities,
        "relationships": relationships,
        "important_leads": leads,
        "failed_stages": failed_stages,
        "counts": {
            "evidence": len(evidence),
            "entities": len(entities),
            "transactions": len(transactions),
            "messages": len(messages),
            "social_profiles": len(social_profiles),
            "technical_indicators": len(technical_indicators),
            "timeline_events": len(timeline),
            "similarity_links": len(similarities),
            "relationships": len(relationships),
            "failed_stages": len(failed_stages),
        },
    }


def _build_leads(transactions: List[Dict[str, Any]], similarities: List[Dict[str, Any]], evidence: List[Dict[str, Any]], messages: List[Dict[str, Any]] | None = None, social_profiles: List[Dict[str, Any]] | None = None, technical_indicators: List[Dict[str, Any]] | None = None) -> List[str]:
    leads: List[str] = []
    accounts = _unique([t.get("account_no") or t.get("receiver_account") for t in transactions])
    utrs = _unique([t.get("utr") for t in transactions])
    if accounts:
        leads.append("Verify/freeze beneficiary account(s): " + ", ".join(accounts[:20]))
    if utrs:
        leads.append("Use UTR/reference values for bank/API reconciliation: " + ", ".join(utrs[:20]))
    if similarities:
        leads.append(f"Review {len(similarities)} evidence similarity link(s), especially financial_link/near_duplicate edges.")
    if messages:
        leads.append(f"Review {len(messages)} structured message(s) for payment requests, threats, OTP requests, phishing links and platform handles.")
    if social_profiles:
        leads.append(f"Verify {len(social_profiles)} social profile/handle record(s) against screenshots/platform evidence.")
    if technical_indicators:
        leads.append(f"Correlate {len(technical_indicators)} technical indicator(s), including URLs/domains/IP/device/QR/payment indicators where present.")
    ai_ready = [e for e in evidence if str(e.get("status", "")).upper() == "AI_READY"]
    if evidence and len(ai_ready) < len(evidence):
        leads.append("Some evidence is not AI_READY yet; check stage errors before final reporting.")
    if not leads:
        leads.append("No high-confidence financial leads stored yet; upload or process financial evidence.")
    return leads


def _unique(values: List[Any]) -> List[str]:
    seen, out = set(), []
    for v in values:
        s = str(v or "").strip()
        if s and s.upper() not in seen:
            seen.add(s.upper())
            out.append(s)
    return out


def _money(v: Any) -> str:
    if v in (None, ""):
        return "N/A"
    try:
        return f"₹{float(v):,.2f}"
    except (TypeError, ValueError):
        return str(v)


def build_report(case_uid: str, kind: str = "investigation") -> Optional[str]:
    """Return a Markdown report of the requested kind, or None if not found."""
    data = build_export_data(case_uid, kind)
    if data is None:
        return None
    case = data["case"]
    sections: List[str] = [
        f"# Investigation Report - {case['title']}",
        "",
        f"- Case UID: {case['uid']}",
        f"- Reference No: {case['reference_no'] or 'N/A'}",
        f"- Status: {case['status']}",
        f"- Generated: {case['generated_at']} (UTC)",
        "",
        "## Executive Summary",
        f"- Evidence items: {data['counts']['evidence']}",
        f"- AI-ready / processed transaction rows: {data['counts']['transactions']}",
        f"- Communication/message records: {data['counts']['messages']}",
        f"- Social profile/handle records: {data['counts']['social_profiles']}",
        f"- Technical indicators: {data['counts']['technical_indicators']}",
        f"- Extracted entities: {data['counts']['entities']}",
        f"- Timeline events: {data['counts']['timeline_events']}",
        f"- Similar evidence links: {data['counts']['similarity_links']}",
        "",
    ]

    if kind in {"investigation", "evidence"}:
        sections.append("## Evidence")
        if data["evidence"]:
            for r in data["evidence"]:
                sections.append(
                    f"- **Evidence #{r['id']} — {r['original_name']}** ({r['mime']}, {r['size']} bytes) "
                    f"- SHA-256 `{r['sha256']}` - status: {r['status']}"
                )
                summary = (r.get("intelligence") or {}).get("summary")
                if summary:
                    sections.append(f"  - Structured summary: {summary}")
        else:
            sections.append("- No evidence recorded.")
        sections.append("")

    if kind in {"investigation", "entity"}:
        sections.append("## Entities")
        if data["entities"]:
            for r in data["entities"][:200]:
                sections.append(f"- {r['type']}: `{r['value']}` (linked to {r['links']} item(s))")
        else:
            sections.append("- No entities extracted.")
        sections.append("")

    if kind in {"investigation", "transactions", "evidence"}:
        sections.append("## Transactions / Money Trail")
        if data["transactions"]:
            for t in data["transactions"]:
                left = t.get("sender_account") or "Participant/source account"
                right = t.get("receiver_account") or t.get("account_no") or "Beneficiary/account"
                layer = f"Layer {t['layer']}" if t.get("layer") is not None else "Layer N/A"
                sections.append(
                    f"- {layer}: `{left}` → `{right}` | UTR `{t.get('utr') or 'N/A'}` | "
                    f"Amount {_money(t.get('amount'))} | IFSC `{t.get('ifsc') or 'N/A'}` | "
                    f"Bank {t.get('bank') or 'N/A'} | Date {t.get('txn_date') or 'N/A'} | "
                    f"Source Evidence #{t['evidence_id']} ({t.get('original_name') or t.get('source_file')}) {t.get('source_ref') or ''}"
                )
        else:
            sections.append("- No structured transactions extracted.")
        sections.append("")

    if kind in {"investigation", "messages", "evidence"}:
        sections.append("## Communication / Message Trail")
        if data["messages"]:
            for m in data["messages"][:300]:
                flags = ", ".join(m.get("risk_flags") or []) or "none"
                sections.append(
                    f"- {m.get('timestamp') or 'undated'} | {m.get('platform') or 'message'} | "
                    f"{m.get('sender') or 'unknown'}: {str(m.get('message_text') or '')[:220]} "
                    f"| flags: {flags} | Source Evidence #{m['evidence_id']} ({m.get('original_name')}) {m.get('source_ref') or ''}"
                )
        else:
            sections.append("- No structured message records extracted.")
        sections.append("")

    if kind in {"investigation", "social", "evidence"}:
        sections.append("## Social Profiles / Handles")
        if data["social_profiles"]:
            for sp in data["social_profiles"][:300]:
                sections.append(
                    f"- {sp.get('platform') or 'social'}: `{sp.get('username') or 'N/A'}` "
                    f"URL: {sp.get('profile_url') or 'N/A'} | Source Evidence #{sp['evidence_id']} ({sp.get('original_name')})"
                )
        else:
            sections.append("- No social profile/handle records extracted.")
        sections.append("")

    if kind in {"investigation", "technical", "evidence"}:
        sections.append("## Technical / Forensic Indicators")
        if data["technical_indicators"]:
            for ti in data["technical_indicators"][:400]:
                sections.append(
                    f"- {ti.get('type')}: `{ti.get('value')}` | Source Evidence #{ti['evidence_id']} ({ti.get('original_name')}) {ti.get('source_ref') or ''}"
                )
        else:
            sections.append("- No technical indicators extracted.")
        sections.append("")

    if kind in {"investigation", "timeline"}:
        sections.append("## Timeline")
        if data["timeline"]:
            for e in data["timeline"]:
                sections.append(f"- {e['ts'] or 'undated'}: {e['summary']}")
        else:
            sections.append("- No timeline events.")
        sections.append("")

    sections.append("## Similar / Linked Evidence")
    if data["similar_evidence"]:
        for s in data["similar_evidence"][:100]:
            reason_bits = []
            for reason in (s.get("reasons") or [])[:5]:
                vals = reason.get("values")
                if vals:
                    reason_bits.append(f"{reason.get('label', reason.get('type'))}: {', '.join(map(str, vals[:6]))}")
                else:
                    reason_bits.append(str(reason.get("label") or reason.get("type") or reason))
            sections.append(
                f"- Evidence #{s['a_id']} ({s.get('a_name')}) ↔ Evidence #{s['b_id']} ({s.get('b_name')}): "
                f"{round(float(s['score']) * 100, 1)}% `{s['kind']}`"
                + (f" — {'; '.join(reason_bits)}" if reason_bits else "")
            )
    else:
        sections.append("- No evidence similarity links stored.")
    sections.append("")

    sections.append("## Important Leads / Recommended Actions")
    for lead in data["important_leads"]:
        sections.append(f"- {lead}")
    sections.append("")

    if data["failed_stages"]:
        sections.append("## Processing Warnings")
        for f in data["failed_stages"]:
            sections.append(f"- Evidence #{f['evidence_id']} ({f['original_name']}), stage `{f['stage']}`: {f['state']} — {f['detail']}")
        sections.append("")

    return "\n".join(sections).rstrip() + "\n"


def build_json(case_uid: str, kind: str = "investigation") -> Optional[str]:
    data = build_export_data(case_uid, kind)
    if data is None:
        return None
    return json.dumps(data, ensure_ascii=False, indent=2, default=str)


def build_pdf_bytes(case_uid: str, kind: str = "investigation") -> Optional[bytes]:
    """Render the Markdown report into a simple text PDF using PyMuPDF."""
    md = build_report(case_uid, kind)
    if md is None:
        return None
    try:
        import fitz  # PyMuPDF, optional offline dependency
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("PDF export requires PyMuPDF (fitz), already listed in requirements.txt") from exc

    doc = fitz.open()
    margin = 42
    fontsize = 9.5
    line_height = 13
    page = doc.new_page(width=595, height=842)
    y = margin
    for raw_line in md.splitlines():
        line = raw_line.rstrip() or " "
        # Basic wrapping to A4 width.
        words = line.split(" ")
        wrapped: List[str] = []
        current = ""
        for word in words:
            test = (current + " " + word).strip()
            if len(test) > 95:
                wrapped.append(current)
                current = word
            else:
                current = test
        wrapped.append(current)
        for part in wrapped:
            if y > 800:
                page = doc.new_page(width=595, height=842)
                y = margin
            page.insert_text((margin, y), part, fontsize=fontsize, fontname="helv")
            y += line_height
    return doc.tobytes()


def build_docx_bytes(case_uid: str, kind: str = "investigation") -> Optional[bytes]:
    """Render the Markdown report into a basic DOCX using python-docx."""
    md = build_report(case_uid, kind)
    if md is None:
        return None
    try:
        from docx import Document  # optional offline dependency
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("DOCX export requires python-docx, already listed in requirements.txt") from exc

    doc = Document()
    for line in md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
