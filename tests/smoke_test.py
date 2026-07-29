"""End-to-end smoke test for the combined local portal."""
from __future__ import annotations

import io
import json
import logging
import os
import shutil
import sys
import tempfile
import time
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook
from werkzeug.test import Client
from werkzeug.wrappers import Response


RUNTIME = tempfile.mkdtemp(prefix="signal-desk-smoke-")
os.environ["SIGNAL_DESK_DATA_DIR"] = RUNTIME
os.environ.pop("GEMINI_API_KEY", None)
os.environ.pop("GOOGLE_API_KEY", None)
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from portal import app  # noqa: E402


client = Client(app, Response)


def expect(response, status=200):
    assert response.status_code == status, (
        response.status_code,
        response.get_data(as_text=True)[:1000],
    )
    return response


def json_body(response):
    return json.loads(response.get_data(as_text=True))


def build_workbook() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Notices"
    sheet.append([
        "Acknowledgement No", "Bank", "Layer", "Account No", "IFSC Code",
        "Transaction Date", "Transaction ID / UTR Number",
        "Transaction Amount", "Reference No", "Remarks",
        "Action Taken by Bank", "Date of Action", "Company Email",
    ])
    sheet.append([
        "ACK-1042", "Atlas Payments", "1", "1234567890", "ATLS0000123",
        "29-07-2026", "TXN-77881234", "25000", "REF-2048",
        "Please review", "Pending", "", "review@atlas.example",
    ])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def run_notice_flow():
    expect(client.get("/"))
    expect(client.get("/notices/"))
    upload = client.post(
        "/notices/api/upload",
        data={"file": (io.BytesIO(build_workbook()), "records.xlsx")},
    )
    expect(upload)
    assert json_body(upload)["ok"] is True

    listing = expect(client.get("/notices/api/records/list"))
    record = json_body(listing)["records"][0]
    updated = client.patch(
        f"/notices/api/records/{record['id']}",
        data=json.dumps({
            "reference_name": "Internal Review",
            "company_email": "review@atlas.example",
        }),
        content_type="application/json",
    )
    expect(updated)

    preview = expect(client.get(
        f"/notices/api/records/preview/{record['id']}?"
        "reference_name=Internal+Review&company_email=review%40atlas.example"
        "&sender_role=No+signature"
    ))
    preview_text = json.dumps(json_body(preview))
    assert "Atlas Payments" in preview_text
    assert "Internal Review" in preview_text

    started = client.post(
        "/notices/api/generate/start",
        data=json.dumps({"sender_role": "No signature", "sender_name": ""}),
        content_type="application/json",
    )
    expect(started)
    job_id = json_body(started)["job_id"]
    job = {}
    deadline = time.time() + 20
    while time.time() < deadline:
        status = expect(client.get(f"/notices/api/generate/status/{job_id}"))
        job = json_body(status)
        if job["state"] in {"ready", "failed"}:
            break
        time.sleep(0.15)
    assert job["state"] == "ready", job
    assert job["generated"] == 1
    assert job["email_drafts"] == 1

    archive_response = expect(client.get(f"/notices/api/generate/download/{job_id}"))
    with zipfile.ZipFile(io.BytesIO(archive_response.data)) as archive:
        names = archive.namelist()
        docx_name = next(name for name in names if name.endswith(".docx"))
        eml_name = next(name for name in names if name.endswith(".eml"))
        assert "delivery_manifest.csv" in names
        generated_docx = Document(io.BytesIO(archive.read(docx_name)))
        generated_text = "\n".join(p.text for p in generated_docx.paragraphs)
        generated_text += "\n" + "\n".join(
            cell.text
            for table in generated_docx.tables
            for row in table.rows
            for cell in row.cells
        )
        assert "Atlas Payments" in generated_text
        assert "Internal Review" in generated_text
        forbidden = tuple(bytes.fromhex(value).decode() for value in (
            "706f6c696365", "6f666669636572",
            "73746174696f6e20686f757365",
            "636f6d6d697373696f6e6572617465",
        ))
        assert not any(term in generated_text.lower() for term in forbidden)
        message = BytesParser(policy=policy.default).parsebytes(archive.read(eml_name))
        assert message["To"] == "review@atlas.example"
        assert any(
            part.get_filename() and part.get_filename().endswith(".docx")
            for part in message.iter_attachments()
        )


def run_investigation_flow():
    expect(client.get("/investigations/"))
    ai_status = json_body(expect(client.get("/investigations/api/ai/status")))
    assert ai_status["default_provider"] == "local"
    assert set(ai_status["providers"]) == {"local", "gemini"}
    assert ai_status["providers"]["gemini"]["available"] is False
    assert ai_status["providers"]["gemini"]["model"] == "gemini-3.1-flash-lite"

    created = client.post(
        "/investigations/api/cases",
        data=json.dumps({"title": "Vendor review", "reference_no": "REV-42"}),
        content_type="application/json",
    )
    expect(created, 201)
    case_uid = json_body(created)["uid"]

    upload = client.post(
        f"/investigations/api/evidence/{case_uid}/upload",
        data={
            "file": (
                io.BytesIO(
                    b"Transaction ID TXN77889900 account 1234567890 "
                    b"email analyst@example.com amount INR 25000"
                ),
                "notes.txt",
            )
        },
    )
    expect(upload, 201)

    items = []
    deadline = time.time() + 20
    while time.time() < deadline:
        items = json_body(expect(client.get(
            f"/investigations/api/evidence/{case_uid}"
        )))
        if items and str(items[0].get("status", "")).lower() not in {
            "pending", "processing", "queued", "extracting", "structuring", "indexing"
        }:
            break
        time.sleep(0.2)
    assert items, "evidence listing was empty"

    expect(client.get(f"/investigations/api/evidence/{case_uid}/entities"))
    search = expect(client.get(
        f"/investigations/api/evidence/{case_uid}/search?q=analyst"
    ))
    assert isinstance(json_body(search), list)
    fallback = json_body(expect(client.post(
        f"/investigations/api/ai/{case_uid}/ask",
        data=json.dumps({
            "query": "Develop a plausible hypothesis and rank its uncertainties.",
            "mode": "smart",
            "provider": "gemini",
        }),
        content_type="application/json",
    )))
    assert fallback["requested_provider"] == "gemini"
    assert fallback["provider"] == "deterministic"
    assert "GEMINI_API_KEY" in fallback["warning"]

    with (
        patch(
            "investigation_app.services.ai_service.GeminiAdapter.is_available",
            return_value=(True, "gemini-3.1-flash-lite is configured"),
        ),
        patch(
            "investigation_app.services.ai_service.GeminiAdapter.generate",
            return_value="Mocked Gemini grounded analysis.",
        ) as generate,
    ):
        gemini_answer = json_body(expect(client.post(
            f"/investigations/api/ai/{case_uid}/ask",
            data=json.dumps({
                "query": "Develop a plausible hypothesis and rank its uncertainties.",
                "mode": "deep",
                "provider": "gemini",
            }),
            content_type="application/json",
        )))
        assert gemini_answer["provider"] == "gemini"
        assert gemini_answer["mode"] == "deep"
        assert gemini_answer["answer"] == "Mocked Gemini grounded analysis."
        sent_context = generate.call_args.kwargs["context"]
        assert sent_context
        assert sent_context[0].startswith("DATABASE-GROUNDED CASE FACTS:")

    report = expect(client.get(
        f"/investigations/api/ai/{case_uid}/report?kind=investigation"
    ))
    report_text = report.get_data(as_text=True)
    assert "Vendor review" in report_text
    assert "Reference No" in report_text
    expect(client.get(
        f"/investigations/api/ai/{case_uid}/report?"
        "kind=investigation&format=docx&download=1"
    ))


if __name__ == "__main__":
    try:
        run_notice_flow()
        run_investigation_flow()
        print("Signal Desk smoke test passed")
    finally:
        from investigation_app.services import jobs

        jobs._shutdown()
        logging.shutdown()
        shutil.rmtree(RUNTIME, ignore_errors=True)
