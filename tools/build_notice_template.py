"""Build the neutral, one-page Notice Studio DOCX template."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "apps" / "notice-studio" / "notice_template.docx"
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
ACCENT = RGBColor(13, 148, 136)


def set_font(run, *, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_text(doc, text="", *, size=11, bold=False, color=INK, after=6,
             before=0, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    set_font(
        paragraph.add_run(text),
        size=size,
        bold=bold,
        color=color,
        italic=italic,
    )
    return paragraph


def add_label_value(doc, label, value, *, after=3):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    set_font(paragraph.add_run(f"{label}: "), bold=True, color=MUTED)
    set_font(paragraph.add_run(value), color=INK)
    return paragraph


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, RGBColor(46, 116, 181)),
        ("Heading 2", 13, 12, 6, RGBColor(46, 116, 181)),
        ("Heading 3", 12, 8, 4, RGBColor(31, 77, 120)),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    document.core_properties.title = "Neutral Information Notice Template"
    document.core_properties.subject = "Email-ready administrative notice"
    document.core_properties.author = "Signal Desk"
    document.core_properties.last_modified_by = "Signal Desk"
    document.core_properties.keywords = "notice, information request, local"

    add_text(
        document,
        "INFORMATION NOTICE",
        size=10,
        bold=True,
        color=ACCENT,
        after=4,
    )
    add_text(
        document,
        "Transaction information request",
        size=22,
        bold=True,
        color=INK,
        after=18,
    )
    add_text(
        document,
        "{{date}}",
        size=10,
        color=MUTED,
        after=14,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    add_label_value(document, "To", "{{bank}}")
    add_label_value(document, "Email", "{{company_email}}")
    add_label_value(document, "Subject", "{{subject}}", after=14)

    add_text(document, "Hello {{bank}} team,", after=10)
    add_text(
        document,
        "Please review the transaction details below and provide the available "
        "information or the appropriate contact for follow-up.",
        after=10,
    )
    add_label_value(document, "Reference name", "{{reference_name}}")
    add_label_value(
        document,
        "Reference number",
        "{{reference_no}} {{acknowledgement_no}}",
        after=10,
    )

    add_text(document, "{{account_table}}", after=10)
    add_text(
        document,
        "Please include the reference number in your reply. This is an "
        "administrative information request and is not a legal order.",
        after=16,
    )
    add_text(document, "Regards,", after=6)
    add_text(document, "{{sender_name}}", bold=True, after=2)
    add_text(document, "{{sender_role}}", color=MUTED, after=0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_font(
        footer.add_run("Generated locally by Signal Desk"),
        size=8,
        color=RGBColor(148, 163, 184),
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
