"""Structural fallback audit for the neutral notice template and filled output."""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "apps" / "notice-studio" / "notice_template.docx"
SAMPLE = ROOT / "work" / "qa" / "notice-sample.docx"
FORBIDDEN = tuple(bytes.fromhex(value).decode() for value in (
    "706f6c696365",
    "6f666669636572",
    "636f6d6d697373696f6e6572617465",
    "73746174696f6e20686f757365",
    "6e617368696b",
    "6d61686172617368747261",
))


def all_text(document: Document) -> str:
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def package_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        )


def audit():
    template = Document(TEMPLATE)
    sample = Document(SAMPLE)
    assert len(template.sections) == 1
    section = template.sections[0]
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.page_height.inches, 2) == 11.00
    assert all(
        round(value.inches, 2) == 1.00
        for value in (
            section.top_margin, section.right_margin,
            section.bottom_margin, section.left_margin,
        )
    )

    template_text = all_text(template)
    for token in (
        "{{date}}", "{{bank}}", "{{company_email}}", "{{subject}}",
        "{{reference_name}}", "{{reference_no}}", "{{account_table}}",
        "{{sender_name}}", "{{sender_role}}",
    ):
        assert token in template_text, token

    sample_text = all_text(sample)
    assert "{{" not in sample_text and "}}" not in sample_text
    assert "Atlas Payments" in sample_text
    assert "Internal Review" in sample_text
    assert not any(term in sample_text.lower() for term in FORBIDDEN)
    assert len(sample.tables) == 1

    table = sample.tables[0]
    assert len(table.columns) == 5
    tbl_pr = table._tbl.tblPr
    assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == "9360"
    assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == "120"
    grid = [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    assert grid == [1944, 1800, 2232, 1656, 1728]
    assert sum(grid) == 9360

    for path in (TEMPLATE, SAMPLE):
        raw = package_text(path).lower()
        assert not any(term in raw for term in FORBIDDEN), path

    assert template.core_properties.author == "Signal Desk"
    assert template.core_properties.last_modified_by == "Signal Desk"
    print("Notice DOCX structural audit passed")


if __name__ == "__main__":
    try:
        audit()
    except AssertionError as error:
        print(f"Notice DOCX structural audit failed: {error}", file=sys.stderr)
        raise
